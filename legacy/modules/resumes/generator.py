'''
Author:     Sai Vignesh Golla
LinkedIn:   https://www.linkedin.com/in/saivigneshgolla/

Copyright (C) 2024 Sai Vignesh Golla

License:    GNU Affero General Public License
            https://www.gnu.org/licenses/agpl-3.0.en.html
            
GitHub:     https://github.com/GodsScion/Auto_job_applier_linkedIn

'''


import docx
from fpdf import FPDF

def create_resume_docx(user_details, summary, experience, projects, skills, certificates):
    # Create a docx file
    doc = docx.Document()

    # Add user details
    doc.add_heading(user_details['name'], 0)
    doc.add_paragraph(user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'])

    # Add summary
    doc.add_heading('Summary', 1)
    doc.add_paragraph(summary)

    # Add experience
    doc.add_heading('Experience', 1)
    for experience_item in experience:
        doc.add_heading(experience_item['company'], 2)
        doc.add_paragraph(experience_item['role'] + ' | ' + experience_item['dates'])
        doc.add_paragraph(experience_item['achievements'])

    # Add projects
    doc.add_heading('Projects', 1)
    for project in projects:
        doc.add_heading(project['name'], 2)
        doc.add_paragraph(project['description'] + ' | ' + project['technologies'])

    # Add skills
    doc.add_heading('Skills', 1)
    doc.add_paragraph(', '.join(skills))

    # Add certificates
    doc.add_heading('Certificates', 1)
    for certificate in certificates:
        doc.add_heading(certificate['name'], 2)
        doc.add_paragraph(certificate['description'])

    # Save docx file
    doc.save('resume.docx')

    # Create a pdf file
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)

    # Add user details
    pdf.cell(0, 10, user_details['name'], 0, 1, 'C')
    pdf.cell(0, 10, user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'], 0, 1, 'L')

    # Add summary
    pdf.cell(0, 10, 'Summary', 0, 1, 'L')
    pdf.multi_cell(0, 10, summary)

    # Add experience
    pdf.cell(0, 10, 'Experience', 0, 1, 'L')
    for experience_item in experience:
        pdf.cell(0, 10, experience_item['company'], 0, 1, 'L')
        pdf.cell(0, 10, experience_item['role'] + ' | ' + experience_item['dates'], 0, 1, 'L')
        pdf.multi_cell(0, 10, experience_item['achievements'])

    # Add projects
    pdf.cell(0, 10, 'Projects', 0, 1, 'L')
    for project in projects:
        pdf.cell(0, 10, project['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, project['description'] + ' | ' + project['technologies'])

    # Add skills
    pdf.cell(0, 10, 'Skills', 0, 1, 'L')
    pdf.multi_cell(0, 10, ', '.join(skills))

    # Add certificates
    pdf.cell(0, 10, 'Certificates', 0, 1, 'L')
    for certificate in certificates:
        pdf.cell(0, 10, certificate['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, certificate['description'])

    # Save pdf file
    pdf.output('resume.pdf', 'F')